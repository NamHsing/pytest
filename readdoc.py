# -*- coding: cp936 -*-
import docx
document =docx.Document(r'E:\工作文档\20171101调研资料\核心厂商客户化改造20171110.docx')
#读取段落内容
docText='\n\n'.join([paragraph.text.encode('cp936') for paragraph in document.paragraphs])
#读取表格内容
exlText= r' '.join(
[cell.text
    for table in document.tables
        for row in table.rows
            for cell in row.cells
])
print exlText
print docText
